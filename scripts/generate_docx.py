#!/usr/bin/env python3
"""生成比赛作品介绍文档 DOCX。"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- 全局样式 --
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x1A)
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x44)
    return p


def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x55)
    return p


def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    return p


def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()  # spacer
    return table


# ============ 正文 ============

h1('艺树（ArTree）—— 让每一幅画都成为一个世界')

para('平台：HarmonyOS 6.0+', bold=True)
para('作品形态：鸿蒙原生 3D 艺术馆应用（Demo）')
para('目标版本：HarmonyOS 6.0 / 7.0 beta')
doc.add_paragraph()

# ---- 一、项目概述 ----
h1('一、项目概述')

h2('1.1 我们做什么')

para('艺树（ArTree）是一个基于 HarmonyOS ArkGraphics 3D 的沉浸式线上艺术馆平台。用户以第一人称视角在 3D 展厅中自由漫游，走近画作时会触发 AI 语音深度讲解；每幅画都包含一个可"走入"的微型 3D 世界（画中世界），将平面的艺术欣赏拓展为空间化的沉浸体验。')

para('同时，艺树内置了基于真实艺术史悬案的解谜推理系统——以 1934 年根特祭坛画《公正法官》面板失窃事件为蓝本，玩家需要在 18 幅名画中收集线索、推理锁定藏匿画作——让艺术学习变成一场侦探游戏。')

h2('1.2 我们要解决的问题')

add_table(
    ['痛点', '艺树的解法'],
    [
        ['美术馆资源集中线下，地域/时间/成本门槛高', '3D 线上展厅，随时随地"走进"美术馆'],
        ['传统观展走马观花，缺乏深度互动', 'AI 语音导览 + 画中世界沉浸 + 解谜驱动探索'],
        ['新兴艺术家缺乏展示渠道和收入来源', '支持上传作品/场景模型 + 虚拟门票商业闭环（规划中）'],
        ['公众对艺术史感到枯燥、距离感强', '以真实艺术史悬案为剧本，将学术知识转化为游戏叙事'],
    ]
)

h2('1.3 名字的由来与理念')

para('"艺树"两个字的读音，和"艺术"几乎一模一样——这是一个谐音梗，但它不止于此。', bold=True)

para('一棵树，从一颗种子开始，扎根、生长、开枝散叶，最终亭亭如盖。我们希望艺术也是这样——不是高高挂在墙上、离普通人很远的东西，而是像树一样，有根、有生命、可以生长到每一个角落。')

para('艺树的核心理念是"让每一幅画都成为一个世界"（Every Painting, A World to Enter.）。我们不只是一个 3D 画展应用。')
bullet('深度而非表面：每幅画包含可走入的 3D 世界，象征艺术不是平面的——它有深度、有上下文、有隐藏的维度')
bullet('发现而非消费：解谜机制让观者主动参与艺术史探索，而非被动滑屏浏览')
bullet('民主而非精英：任何人都可以上传作品、策展、分享——打破画廊的门槛')
bullet('桥梁而非壁垒：西洋馆与中国馆并列，展示不同文明对"美"的共同追求')

h2('1.4 面向的用户群体')

para('艺树想服务三群人：', bold=True)
bullet('第一，艺术爱好者（18-45 岁）：想去美术馆但没有时间、距离太远、或者门票太贵的人——他们需要一个随时随地能"走进去"的线上艺术空间')
bullet('第二，学生群体：觉得艺术史枯燥、记不住画家和年代——他们需要一种更好玩、更沉浸的学习方式，把知识变成一场侦探游戏')
bullet('第三，新兴艺术家：有作品、有创意，但没有地方展示、没有渠道接触到观众、没有方式获得收入——他们需要一个低门槛的线上展览平台和商业变现路径')

para('这三群人有一个共同点：他们都被传统艺术体系挡在了外面。艺树要做的，就是把这个门槛拆掉。')

# ---- 二、创新场景设计 ----
h1('二、创新场景设计')

h2('2.1 核心创新点')

h3('① 画中世界（Painting World）——从"看画"到"入画"')
para('走近一幅画作后，点击"走入画中"即可进入画作内部的 3D 微缩世界。每幅画拥有独特的配色、地形和氛围：蒙娜丽莎的托斯卡纳丘陵、莫奈日出的勒阿弗尔港晨雾、弗里德里希的云海山巅。基于深度估计算法从 2D 画作中生成 3D 浮雕地形，让观者真正"走入"艺术家的视角。画中世界支持重力物理和跳跃，增强了沉浸式的空间体验。')

h3('② 侦探解谜式艺术学习')
para('以真实艺术史悬案为蓝本——1934 年 Van Eyck《根特祭坛画》中《公正法官》面板失窃，至今未寻回。18 幅画作中散布 6 条"窃语"线索，每条线索对应一个艺术史知识点（如"十五世纪"、"弗拉芒油画技法"、"西西里画派"），集齐后通过逻辑推理唯一锁定目标画作。将枯燥的艺术史知识转化为沉浸式侦探体验。')

h3('③ 双馆叙事架构')
para('"西洋馆"与"中国馆"并列，不将东西方艺术对立，而是让观者在两个空间中自由穿行，感受不同文明对"美"的追求。西洋馆以"根特悬案"为线，中国馆以"兰亭序之谜"为线——两段真实的艺术史失踪案，跨越东西方，共同构成艺树的叙事骨架。')

h3('④ AI 语音深度导览')
para('基于鸿蒙 Core Speech Kit，每幅画拥有专属 AI 语音旁白。不是简单朗读画名和年代，而是讲述这幅画背后的故事、技法创新和人文内涵。字幕与语音实时同步，支持静音模式。语音参数（语速/音高）可调，营造神秘古典的夜间艺术馆氛围。')

h3('⑤ 完整的沉浸式氛围')
para('内置文艺复兴时期古典背景音乐（可开关），4 方向动态光照系统（主光+冷补光+暖侧光+天顶补光），PBR 真实材质，AABB 碰撞检测，矢量小地图实时导航。模糊透明美学 UI（毛玻璃模态面板、半透明 HUD 组件）符合 HarmonyOS 设计规范。')

h2('2.2 与同类产品的差异')

add_table(
    ['对比维度', '传统线上画展', 'Google Arts & Culture', '艺树（ArTree）'],
    [
        ['交互方式', '点击翻页/全景', '缩放浏览', '第一人称自由漫游'],
        ['深度理解', '文字说明', '文字+短视频', 'AI语音+画中世界+解谜'],
        ['创作参与', '无', '无', '支持上传作品/场景'],
        ['商业模式', '免费', '免费', '虚拟门票+艺术家分成'],
        ['平台', 'Web/VR', 'Web/App', '鸿蒙原生（全场景协同）'],
    ]
)

# ---- 三、用户与市场分析 ----
h1('三、用户与市场分析')

h2('3.1 核心用户群')

bullet('艺术爱好者（18-45 岁）：希望深入了解艺术但缺乏线下观展条件的人群')
bullet('学生群体：需要生动有趣的艺术史学习方式')
bullet('新兴艺术家：缺少展示空间和变现渠道的创作者')
bullet('策展人/画廊：需要低成本的线上展览解决方案')

h2('3.2 市场数据')

bullet('全球线上艺术市场 2024 年规模约 118 亿美元（Statista），预计 2030 年达 250 亿美元')
bullet('中国数字艺术消费用户已超 3 亿（QuestMobile 2025）')
bullet('线下美术馆单次特展成本 50-500 万元，线上 3D 展厅建设成本仅为其 1%-5%')
bullet('鸿蒙生态设备超过 10 亿台，覆盖手机、平板、智慧屏、车机等')
bullet('大英博物馆、卢浮宫在疫情期间线上收入增长 300%+，"线上观展+数字门票"模式已被验证')

h2('3.3 商业闭环（规划中）')
para('艺术家上传作品 → 3D 展厅自动生成 → 观众购买虚拟门票 → 平台抽成 15-30% → 艺术家获得门票收入 + 曝光 → 优质展览吸引更多观众 → 正向循环。')

# ---- 四、HarmonyOS 技术集成 ----
h1('四、HarmonyOS 能力技术集成方案')

h2('4.1 使用的鸿蒙开放能力')

add_table(
    ['鸿蒙能力', '用途', '创新主题'],
    [
        ['ArkGraphics 3D', '3D 展厅渲染、PBR 材质、实时光影、第一人称相机', '3D 空间化'],
        ['Core Speech Kit', 'AI 语音旁白（TTS 文本转语音）', 'AI 智能化体验'],
        ['ArkUI', '模糊透明美学 UI、矢量小地图、虚拟摇杆', 'UX 美学'],
        ['ArkWeb', 'Three.js 画中世界渲染', '3D 空间化'],
        ['Media Kit（AVPlayer）', '文艺复兴背景音乐循环播放', '沉浸体验'],
        ['Core Vision Kit（规划）', '画作识别、自动标签、内容审核', 'AI 智能化体验'],
        ['Account Kit（规划）', '艺术家身份认证', '安全隐私保护'],
        ['Distributed Data（规划）', '多设备协同观展', '全场景一体协同'],
    ]
)

h2('4.2 核心技术实现')

h3('3D 渲染管线')
para('glTF 2.0 场景模型（离线 Python 生成）→ ArkGraphics 3D Scene.load() → PBR 材质 + 4 方向光照 → Component3D 嵌入 ArkUI → 60fps 实时渲染')

h3('画中世界生成')
para('2D 画作纹理 → 深度估计模型（离线）→ 64×64 深度图 → 顶点位移生成 3D 浮雕地形 → Three.js WebView 渲染 → 重力 + 跳跃物理 + 地面碰撞')

h3('AI 语音管线')
para('遇到画作 → speak(title+description) → Core Speech Kit TTS Engine → 字幕回调 → ArkUI Text 实时显示 → 支持静音模式（仅显示字幕）')

h3('解谜推理引擎')
para('收集线索 → 按艺术史分类排除画作 → 候选集逐步缩小（18→5→2→1）→ 集齐 6 条线索唯一锁定目标')

# ---- 五、安全隐私保护 ----
h1('五、安全隐私保护方案')

bullet('艺术家上传作品存储于鸿蒙应用沙箱，外部应用无法访问')
bullet('用户观展记录仅保存在本地，不上传服务器')
bullet('虚拟门票交易通过 HarmonyOS IAP Kit 完成，不接触用户支付信息')
bullet('上传作品经 Core Vision Kit 自动内容审核（规划中）')
bullet('展览内容可设置访问权限：公开 / 仅门票持有者 / 仅受邀者')
bullet('数字水印嵌入上传作品，防止未经授权的二次传播')

# ---- 六、全场景协同 ----
h1('六、全场景一体协同方案')
para('（当前版本已支持 tablet/2in1/phone 三种设备类型，以下为规划中的分布式能力）')
bullet('手机扫码进入展览 → 就近画作 → 轻触"投屏"→ 平板/智慧屏自动接续展示高清细节')
bullet('手机变为语音导览遥控器 + 字幕显示器')
bullet('手表同步震动提醒"前方有隐藏线索"')
bullet('多人同时在线观展，看到彼此位置和视角（虚拟化身）')
bullet('策展人在平板上编辑展览布局，观者手机端实时更新')
bullet('基于鸿蒙分布式软总线，设备间延迟 < 20ms')

# ---- 七、开源与技术共享 ----
h1('七、可开源共享的技术/组件')

add_table(
    ['组件', '说明', '许可证'],
    [
        ['glTF Scene Generator', 'Python 程序化 3D 展厅生成器', 'Apache 2.0'],
        ['Painting World Builder', '2D→3D 画中世界生成管线（深度估计+glTF构建）', 'Apache 2.0'],
        ['SpeechEngine', '鸿蒙 Core Speech Kit 封装（语音+字幕同步）', 'Apache 2.0'],
        ['Detective Engine', '基于逻辑约束的推理解谜框架', 'Apache 2.0'],
        ['Exhibition JSON Format', '展览数据交换格式规范', 'CC0'],
    ]
)

para('开源场景化能力：')
bullet('程序化展厅生成：输入画作列表 + 布局参数，自动生成 glTF 2.0 场景')
bullet('画作纹理管线：批量下载公有领域画作 → 压缩 → 嵌入 glTF')
bullet('深度估计→浮雕地形：基于深度学习，将任意 2D 画作转为 3D 地形')

# ---- 八、技术架构 ----
h1('八、整体技术架构')

para('应用分层架构：', bold=True)
bullet('ArkUI 层：Index（首页）+ DavidScene（西洋馆）+ 中国馆（规划）+ Components（Minimap、Joystick、ExhibitCard、ObservationModal、NotebookModal、SettingsModal）')
bullet('业务逻辑层：SpeechEngine（TTS+字幕同步）、BgmEngine（背景音乐）、DetectiveEngine（线索推理）')
bullet('鸿蒙系统能力层：ArkGraphics3D（3D 渲染）、Core Speech Kit（TTS）、Media Kit（音频）、ArkWeb（WebView）、Account Kit（认证）、Distributed Data（协同）')
bullet('离线资产生成层：Python Scripts（generate_gallery.py、add_paintings.py、depth_estimate.py）')

# ---- 九、当前功能清单 ----
h1('九、当前版本功能清单（Phase 1）')

bullet('西洋馆（大卫馆）第一人称 3D 自由漫游')
bullet('18 幅世界名画展示（从文艺复兴到印象派，全部公有领域、无宗教裸露题材）')
bullet('AI 语音深度导览 + 实时字幕（Core Speech Kit TTS）')
bullet('根特悬案推理解谜系统（6 条线索 → 18→5→2→1 推理链）')
bullet('画中世界 3D 微缩场景（WebView 路径，含重力物理+双击跳跃）')
bullet('文艺复兴背景音乐 + 语音/音乐独立开关')
bullet('矢量小地图实时导航')
bullet('AABB 碰撞检测')
bullet('模糊透明美学 UI（毛玻璃模态面板）')
bullet('中国馆数据架构（12 幅传世名画 + 兰亭序之谜线索系统，待 3D 场景接入）')

# ---- 十、未来规划 ----
h1('十、未来规划与商业路径')

add_table(
    ['阶段', '时间', '内容'],
    [
        ['Phase 1 · 已完成', '2026.5', '西洋馆 3D 漫游 + 18 幅名画 + AI 语音 + 解谜系统 + 画中世界 + 背景音乐'],
        ['Phase 2 · 近期', '1-2 月', '中国馆 3D 场景接入 + 用户上传作品 + Core Vision Kit 画作识别'],
        ['Phase 3 · 中期', '3-6 月', '虚拟门票系统（IAP Kit）+ 艺术家收益分成 + 多设备协同观展 + 手机/手表适配'],
        ['Phase 4 · 远期', '6-12 月', '用户自建展厅（上传 3D 模型）+ 社交观展 + AI 策展助手 + 全球化多语言'],
    ]
)

# ---- 十一、社会价值 ----
h1('十一、社会价值与愿景')

para('艺树（ArTree）的使命不是替代线下美术馆，而是让那些因为地理距离、经济条件或身体限制而无法亲临美术馆的人，也能走进艺术的殿堂。')
para('')
para('我们相信：艺术不应有门槛，美应该在每一个人触手可及的地方。', bold=True)
para('')
para('通过技术手段降低艺术教育的成本、通过商业模式支持艺术家的创作、通过沉浸式体验让更多人爱上艺术——这就是艺树存在的意义。')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('—— 让每一幅画都成为一个世界 ——')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x77)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('项目名称：艺树（ArTree）\n参赛作者：SowrJam\n提交日期：2026 年 5 月 15 日')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x99)

# -- 保存 --
out_path = os.path.join('docs', '艺树ArTree_作品介绍文档_v2.docx')
os.makedirs('docs', exist_ok=True)
doc.save(out_path)
print(f'Saved: {out_path}')
